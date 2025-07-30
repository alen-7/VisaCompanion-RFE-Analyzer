from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# The code above is incorrectly indented and not inside any function.
# If you want to keep this logic, move it inside a function (e.g., add_auto_fix_suggestions).
# Otherwise, remove it if not needed.

# Example: Move inside a function (assuming you want to keep it):

def add_auto_fix_suggestions(doc, risks):
    """Adds before/after table for suggested fixes."""
    doc.add_paragraph("Auto-Fix Suggestions", style='Heading1')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'LightShading-Accent2'
    headers = ["Category", "Original Issue", "Fixed Text", "Example"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    # Create a table with before/after examples
    for criterion, issue, severity, fixed_text in risks:  # <-- Now unpacking 4 values
        row = table.add_row().cells
        row[0].text = criterion
        row[1].text = issue
        row[2].text = fixed_text

        # Add click-to-copy fixed text
        p = row[2].paragraphs[0]
        p.add_run("Original: ").bold = True
        p.add_run(issue.split(":")[-1].strip())

        p = row[3].paragraphs[0]
        p.add_run("Fixed: ").bold = True
        p.add_run(fixed_text)
        

def add_uscis_perspective(doc, risks):
    """Adds mock USCIS officer commentary"""
    doc.add_heading("Adjudicator's Perspective", level=2)
    
    # Custom comments based on found risks
    if any(r[2] == "High" for r in risks):
        doc.add_paragraph(
            "As a USCIS officer reviewing this petition, I would question:",
            style='IntenseQuote'
        )
        doc.add_paragraph(
            "🔴 Substantiation Issues: Several claims lack independent verification "
            "(e.g., no patent numbers for 'innovative algorithms'). RFE likely required.",
            style='ListBullet'
        )
    else:
        doc.add_paragraph(
            "🟡 Generally meets criteria but needs strengthening in:",
            style='IntenseQuote'
        )
    
    # Always include these standard notes
    standard_notes = [
        "Compare to Matter of [XXXX] (AAO precedent case)",
        "Verify if publications meet 6.2.3(b) criteria",
        "Check for objective evidence under Kazarian framework"
    ]
    for note in standard_notes:
        doc.add_paragraph(note, style='ListBullet')
# -------------------------------

def generate_report(risks, output_path):
    doc = Document()
    
    # Add professional header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_run = header_para.add_run()
    try:
        header_run.add_picture('logo.png', width=Inches(1.5))  # Add your logo
    except:
        pass  # Skip if logo not found
    
    # Title
    title = doc.add_paragraph("VisaCompanion RFE Risk Report", style='Title')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Executive summary
    doc.add_paragraph("Executive Summary", style='Heading1')
    high_risks = len([r for r in risks if r[2] == "High"])
    doc.add_paragraph(
        f"This petition review identified {len(risks)} potential RFE triggers "
        f"({high_risks} critical, {len(risks)-high_risks} recommended improvements).",
        style='IntenseQuote'
    )
    
    # Detailed findings
    doc.add_paragraph("Detailed Findings", style='Heading1')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'LightShading-Accent1'
    
    # Table headers
    headers = ["Category", "Issue", "Severity", "USCIS Policy Reference"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # Policy references
    policy_refs = {
        "Original Contributions": "USCIS PM 6.2.1",
        "Publications": "USCIS PM 6.2.3", 
        "Critical Role": "USCIS PM 6.2.5",
        "Subjective Language": "AAO Decision #123456"
    }
    
    for criterion, issue, severity in risks:
        row = table.add_row().cells
        row[0].text = criterion
        row[1].text = issue
        row[2].text = severity
        row[3].text = policy_refs.get(criterion, "General RFE Pattern")
        
        # Color coding
        if severity == "High":
            row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            row[1].paragraphs[0].runs[0].font.bold = True
        elif severity == "Medium":
            row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)
    
    # Improvement templates
    doc.add_paragraph("Suggested Wording Improvements", style='Heading1')
    doc.add_paragraph("Before → After Examples:", style='Heading2')
    improvements = [
        ('"world-class"', '"ACM Fellow (top 1% of CS researchers)"'),
        ('"innovative"', '"Patented algorithm (USPTO #123456)"'),
        ('"led team"', '"Directed 5-person team (org chart Exhibit B)"')
    ]
    for bad, good in improvements:
        doc.add_paragraph(f"{bad} → {good}", style='ListBullet')
    add_uscis_perspective(doc, risks) 
    add_auto_fix_suggestions(doc, risks)

    doc.save(output_path)