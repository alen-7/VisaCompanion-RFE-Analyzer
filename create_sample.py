from docx import Document
import os

# Create directories if they don't exist
os.makedirs("input", exist_ok=True)
os.makedirs("output", exist_ok=True)

# Create a sample petition document
doc = Document()
doc.add_heading('EB-1A Extraordinary Ability Petition', 0)

# Add sections with both strong and weak claims
doc.add_heading('Personal Statement', level=1)
doc.add_paragraph('I am a world-class researcher in artificial intelligence with unparalleled expertise.')

doc.add_heading('Original Contributions', level=1)
doc.add_paragraph('Developed innovative algorithms that improved efficiency by 40%.')

doc.add_heading('Publications', level=1)
doc.add_paragraph('Published 3 papers in conference proceedings (no journal citations provided).')

doc.add_heading('Critical Role', level=1)
doc.add_paragraph('Led a team of 5 researchers at XYZ Corporation.')

# Save the document
doc.save('input/sample_petition.docx')
print("Sample petition created at: input/sample_petition.docx")